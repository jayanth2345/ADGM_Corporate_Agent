from docx import Document

def inject_comments(input_path, output_path, issues):
    doc = Document(input_path)
    for para in doc.paragraphs:
        text = para.text.lower()
        for issue in issues:
            section = issue.get("section", "").lower()
            if section in text or issue["issue"][:15].lower() in text:
                run = para.add_run()
                run.text = f" [COMMENT: {issue['issue']} Suggestion: {issue['suggestion']} (Ref: ADGM Rule)] "
                run.font.color.rgb = (255, 0, 0)
                break
    doc.save(output_path)