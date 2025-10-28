# create_samples.py
from docx import Document
import os

# Create sample_input/
os.makedirs("sample_input", exist_ok=True)

# Sample: Articles_of_Association.docx
doc = Document()
doc.add_heading("Articles of Association", 0)
doc.add_paragraph("CLAUSE 3.1: Jurisdiction\nThis agreement shall be governed by the laws of the UAE and subject to the jurisdiction of Dubai Courts.")
doc.add_paragraph("CLAUSE 4.2: Share Capital\nThe company shall have a share capital of AED 100,000 divided into 10,000 shares.")
doc.save("sample_input/Articles_of_Association.docx")

print("✅ Created sample_input/Articles_of_Association.docx")